from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Cm
from django.conf import settings
import os


def generate_docx(context):
    os.path.join(settings.BASE_DIR, 'form', 'templates')

    print(os.path.join(settings.BASE_DIR,'form', 'templates', 'template.docx'))
    doc = DocxTemplate(os.path.join(settings.BASE_DIR,'form', 'templates', 'template.docx'))
    # print(os.path.abspath("templates/form/template.docx")
    doc.render(context)
    doc.save(os.path.join(settings.BASE_DIR,'static', 'docx_files', 'шаблон-final.docx'))

    document = Document(os.path.join(settings.BASE_DIR,'static', 'docx_files', 'шаблон-final.docx'))

    # table = document.add_table(rows=2, cols=11,style = 'Table Normal')

    # styles = document.styles
    # style = styles[WD_STYLE.TABLE_MEDIUM_GRID_1]
    # document.styles = None
    # table.style = 'Table Grid'

    table = document.tables[1]

    # Список с данными для таблицы
    context_new = [
        'Обеспечение учебно-воспитательного процесса и самообразования населения путём обслуживания читателей и предоставление библиотечного фонда',
        [

            [
                'Опасность недостаточной освещенности в рабочей зоне',
                'Снижение остроты зрения',
                'Потеря трудоспособности, инвалидность, профзаболевание',
                'Т',
                'Проведение регулярных замеров освещенности на рабочих местах (производственный контроль). Использование переносных светильников для освещения рабочей зоны. ',
                '3',
                '2',
                '6',
                'Нет',
                ' '
            ],
            [
                'Опасность недостаточной освещенности в рабочей зоне',
                'Снижение остроты зрения',
                'Потеря трудоспособности, инвалидность, профзаболевание',
                'Т',
                'Проведение регулярных замеров освещенности на рабочих местах (производственный контроль). Использование переносных светильников для освещения рабочей зоны. ',
                '3',
                '2',
                '6',
                'Нет',
                ' '
            ],
            [
                'Опасность недостаточной освещенности в рабочей зоне',
                'Снижение остроты зрения',
                'Потеря трудоспособности, инвалидность, профзаболевание',
                'Т',
                'Проведение регулярных замеров освещенности на рабочих местах (производственный контроль). Использование переносных светильников для освещения рабочей зоны. ',
                '3',
                '2',
                '6',
                'Нет',
                ' '
            ],
        ],
        'Обеспечение учебно-воспитательного процесса и самообразования населения путём обслуживания читателей и предоставление библиотечного фонда',
        [
            [
                'Опасность недостаточной освещенности в рабочей зоне',
                'Снижение остроты зрения',
                'Потеря трудоспособности, инвалидность, профзаболевание',
                'Т',
                'Проведение регулярных замеров освещенности на рабочих местах (производственный контроль). Использование переносных светильников для освещения рабочей зоны. ',
                '3',
                '2',
                '6',
                'Нет',
                ' '
            ],
            [
                'Опасность недостаточной освещенности в рабочей зоне',
                'Снижение остроты зрения',
                'Потеря трудоспособности, инвалидность, профзаболевание',
                'Т',
                'Проведение регулярных замеров освещенности на рабочих местах (производственный контроль). Использование переносных светильников для освещения рабочей зоны. ',
                '3',
                '2',
                '6',
                'Нет',
                ' '
            ],
            [
                'Опасность недостаточной освещенности в рабочей зоне',
                'Снижение остроты зрения',
                'Потеря трудоспособности, инвалидность, профзаболевание',
                'Т',
                'Проведение регулярных замеров освещенности на рабочих местах (производственный контроль). Использование переносных светильников для освещения рабочей зоны. ',
                '3',
                '2',
                '6',
                'Нет',
                ' '
            ],
        ]
    ]

    # Измерение кол-ва строк
    lenth = 0
    print(round(len(context_new) / 2))
    for i in range(round(len(context_new) / 2)):
        lenth += len(context_new[1])
    print(lenth)
    print(context_new[1][0])
    doc.render(context)
    doc.save("шаблон-final.docx")

    sections = document.sections

    # убираем отступы полей
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    # Индекс для объединения строк в столбце №1
    index_cell_start = 2
    index_cell_end = 2

    for i in range(lenth):
        row = table.add_row()

    # Цикл для объединения строк в столбце №1
    for i in range(1, len(context_new), 2):
        index_cell_end += round(len(context_new[i])) - 1
        table.cell(index_cell_start, 0).merge(table.cell(index_cell_end, 0))
        print(index_cell_start)
        print(index_cell_end)
        index_cell_end += 1
        index_cell_start = index_cell_end

    # Цикл для заполнения таблицы
    k = 0
    l = 1
    c = 0
    for i in range(2, index_cell_end):
        row = table.rows[i]
        for j in range(1, 11):
            row.cells[j].text = context_new[l][k][j - 1]
        k += 1
        c += 1
        if k == len(context_new[l]):
            k = 0

        if c == len(context_new[l]):
            l += 2
            c = 0
    counter = 0
    print(len(context_new) / 2)
    for i in range(2, lenth + 2, 3):
        row = table.rows[i]
        row.cells[0].text = context_new[counter]
        counter += 2
        print(1)

    print(table)

    document.save(os.path.join(settings.BASE_DIR,'static', 'docx_files', 'шаблон-final.docx'))

